"""
Script to generate professional Word document for Pattern Learning System
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def create_pattern_learning_doc():
    doc = Document()

    # Set up styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Title Page
    title = doc.add_heading('Warehouse Putaway System', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('Pattern Learning & Recommendation Architecture')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.runs[0]
    subtitle_format.font.size = Pt(16)
    subtitle_format.font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph()

    # Metadata
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run('Document Version: 2.1\n').bold = True
    meta.add_run('Last Updated: February 2026\n')
    meta.add_run('Status: Production Implementation')

    doc.add_page_break()

    # Table of Contents
    doc.add_heading('Table of Contents', 1)
    toc_items = [
        '1. Executive Summary',
        '2. System Architecture',
        '3. Data Foundation',
        '4. Pattern Extraction Process',
        '5. Recommendation Engine',
        '6. Implementation Details',
        '7. Performance Metrics',
        '8. Operational Guidelines',
        'Appendix A: Sample Data',
        'Appendix B: File Structure',
        'Appendix C: API Reference'
    ]
    for item in toc_items:
        p = doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # 1. Executive Summary
    doc.add_heading('1. Executive Summary', 1)

    doc.add_heading('1.1 Overview', 2)
    doc.add_paragraph(
        'The Warehouse Putaway Recommendation System leverages historical transaction data '
        'to provide intelligent storage location recommendations. The system analyzes operator '
        'behavior patterns from 224,081 putaway transactions across 2,409 parts and 83 clients, '
        'spanning 3.7 years of warehouse operations (January 2022 - September 2025).'
    )

    doc.add_heading('1.2 Core Objective', 2)
    doc.add_paragraph(
        'Design and implement an intelligent recommendation engine that suggests optimal '
        'warehouse storage locations based on:'
    )
    objectives = [
        'Historical operator behavior patterns',
        'Real-time location availability',
        'Client-specific storage preferences',
        'Warehouse zone optimization'
    ]
    for obj in objectives:
        doc.add_paragraph(obj, style='List Bullet')

    doc.add_heading('1.3 System Components', 2)

    # Create components table
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Light Grid Accent 1'

    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Component'
    header_cells[1].text = 'Technology'
    header_cells[2].text = 'Purpose'

    # Make header bold
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    # Data rows
    components_data = [
        ('Pattern Store', 'Qdrant Vector DB', 'Historical pattern repository (2,492 documents)'),
        ('Real-Time Database', 'MySQL', 'Current inventory & availability state'),
        ('Pattern Engine', 'Python 3.8+', 'Statistical analysis & aggregation'),
        ('AI Explanation Layer', 'Ollama LLM (llama3.2)', 'Natural language recommendations')
    ]

    for idx, (comp, tech, purpose) in enumerate(components_data, 1):
        row = table.rows[idx]
        row.cells[0].text = comp
        row.cells[1].text = tech
        row.cells[2].text = purpose

    doc.add_paragraph()

    doc.add_heading('1.4 Key Metrics', 2)

    metrics_table = doc.add_table(rows=7, cols=2)
    metrics_table.style = 'Light List Accent 1'

    metrics_data = [
        ('Total Transactions Analyzed', '224,081'),
        ('Unique Parts', '2,409'),
        ('Active Clients', '83'),
        ('Warehouse Locations', '31,416'),
        ('Pattern Confidence Range', '0.12 - 0.98'),
        ('Average Response Time', '< 300ms')
    ]

    for idx, (metric, value) in enumerate(metrics_data):
        row = metrics_table.rows[idx]
        row.cells[0].text = metric
        row.cells[1].text = value
        row.cells[1].paragraphs[0].runs[0].font.bold = True

    doc.add_page_break()

    # 2. System Architecture
    doc.add_heading('2. System Architecture', 1)

    doc.add_heading('2.1 High-Level Architecture', 2)
    doc.add_paragraph(
        'The system follows a three-tier architecture combining historical pattern analysis, '
        'real-time validation, and AI-powered explanations:'
    )

    arch_steps = [
        'User Request: Operator provides Part ID for putaway recommendation',
        'Pattern Retrieval: System queries Qdrant for historical putaway patterns',
        'Availability Check: Real-time query to MySQL verifies location status (FREE/OCCUPIED)',
        'AI Generation: Ollama LLM creates natural language explanation',
        'Final Recommendation: Combined result with location, status, and reasoning'
    ]

    for step in arch_steps:
        doc.add_paragraph(step, style='List Number')

    doc.add_heading('2.2 Data Flow Sequence', 2)

    flow_steps = [
        ('Step 1: Request Initiation', 'recommend_putaway(part_id=600)'),
        ('Step 2: Part Identification', 'SELECT code, description, clientId FROM part WHERE id = 600'),
        ('Step 3: Pattern Retrieval', 'qdrant.retrieve(collection="PartSummary", ids=[600])'),
        ('Step 4: Availability Check', 'SELECT clientId FROM location WHERE code = "TP03D"'),
        ('Step 5: AI Explanation', 'ollama.generate(prompt="Recommend location TP03D...")')
    ]

    for step_name, step_code in flow_steps:
        doc.add_paragraph(step_name).runs[0].font.bold = True
        code_para = doc.add_paragraph(step_code)
        code_para.runs[0].font.name = 'Courier New'
        code_para.runs[0].font.size = Pt(9)

    doc.add_page_break()

    # 3. Data Foundation
    doc.add_heading('3. Data Foundation', 1)

    doc.add_heading('3.1 Source Database Schema', 2)

    doc.add_paragraph('The system relies on four core database tables:')

    tables_info = [
        ('Transaction Table', '224,081 records', 'Stores all putaway operations with location, part, client, and timestamp'),
        ('Part Table', '3,215 records', 'Part catalog with code, description, and client assignment'),
        ('Location Table', '31,416 records', 'Warehouse locations with availability status (clientId NULL = FREE)'),
        ('Client Table', '87 records', 'Client information and metadata')
    ]

    for table_name, count, desc in tables_info:
        p = doc.add_paragraph()
        p.add_run(f'{table_name} ').bold = True
        p.add_run(f'({count}): ')
        p.add_run(desc)

    doc.add_heading('3.2 Location Naming Convention', 2)
    doc.add_paragraph(
        'Warehouse locations follow a structured 4-5 character format. Example: TN52D'
    )

    naming_parts = [
        'T = Zone (Temperature Controlled)',
        'N = Aisle (North section)',
        '52 = Column (Physical column number)',
        'D = Row (4th level from ground, A=ground, F=top)'
    ]

    for part in naming_parts:
        doc.add_paragraph(part, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('Invalid Location Patterns (Filtered Out):').runs[0].font.bold = True

    invalid_patterns = [
        'FLOOR* - Temporary floor storage',
        'REC* - Receiving staging area',
        'ORD* - Order picking virtual locations',
        'TN52DD - Subdivided locations (double letter suffix)'
    ]

    for pattern in invalid_patterns:
        doc.add_paragraph(pattern, style='List Bullet')

    doc.add_page_break()

    # 4. Pattern Extraction Process
    doc.add_heading('4. Pattern Extraction Process', 1)

    doc.add_heading('4.1 Aggregation Pipeline', 2)
    doc.add_paragraph('The pattern extraction follows a six-phase pipeline:')

    phases = [
        ('Phase 1: Location Frequency Analysis',
         'Count occurrences of each location for the part using Python Counter'),
        ('Phase 2: Statistical Calculation',
         'Calculate usage percentages and derive confidence metrics'),
        ('Phase 3: Location Parsing',
         'Extract zone, aisle, column, and row from location codes'),
        ('Phase 4: Pattern Strength Classification',
         'Categorize as STRONG (>50%), MODERATE (30-50%), or WEAK (<30%)'),
        ('Phase 5: Consistency Ratio',
         'Measure predictability: unique_locations / total_putaways'),
        ('Phase 6: Temporal Analysis',
         'Track first/last usage dates and calculate recency status')
    ]

    for phase_name, phase_desc in phases:
        p = doc.add_paragraph()
        p.add_run(f'{phase_name}\n').bold = True
        p.add_run(phase_desc)

    doc.add_heading('4.2 Confidence Score Formula', 2)
    doc.add_paragraph(
        'The confidence score combines multiple factors to produce a reliability metric (0.0 - 1.0):'
    )

    formula_components = [
        'Base Score: Determined by pattern strength (STRONG=0.7, MODERATE=0.5, WEAK=0.3)',
        'Consistency Bonus: Lower consistency ratio adds up to 0.2 points',
        'Recency Factor: Multiplier based on days since last use (1.0 if <30 days, 0.5 if >180 days)',
        'Volume Bonus: More historical data adds confidence (up to 0.1 points)'
    ]

    for component in formula_components:
        doc.add_paragraph(component, style='List Bullet')

    doc.add_paragraph()
    p = doc.add_paragraph('Formula: ')
    p.add_run('(base_score + consistency_bonus + volume_bonus) × recency_factor').font.italic = True

    doc.add_page_break()

    # 5. Recommendation Engine
    doc.add_heading('5. Recommendation Engine', 1)

    doc.add_heading('5.1 Decision Workflow', 2)

    workflow_steps = [
        'STEP 1: Identify Part - Query MySQL for part code, description, and client',
        'STEP 2: Retrieve Patterns - Query Qdrant for historical locations and metrics',
        'STEP 3: Check Availability - For each location, verify real-time FREE/OCCUPIED status',
        'STEP 4: Generate Explanation - Call Ollama LLM with recommendation context',
        'STEP 5: Return Recommendation - Provide location, status, alternatives, and AI summary'
    ]

    for step in workflow_steps:
        doc.add_paragraph(step, style='List Number')

    doc.add_heading('5.2 Location Validation Rules', 2)
    doc.add_paragraph('Locations are filtered using the following rules:')

    validation_rules = [
        'Reject if location code is empty or None',
        'Reject if starts with FLOOR, REC, or ORD (temporary/virtual)',
        'Reject if ends with double letters (subdivided locations like TN52DD)',
        'Accept only valid physical warehouse locations'
    ]

    for rule in validation_rules:
        doc.add_paragraph(rule, style='List Bullet')

    doc.add_heading('5.3 Recommendation Sorting Logic', 2)
    doc.add_paragraph('The system prioritizes recommendations using this algorithm:')

    sorting_steps = [
        'Filter: Remove invalid location codes (FLOOR*, REC*, ORD*, subdivisions)',
        'Filter: Remove occupied locations (query MySQL for clientId IS NOT NULL)',
        'Sort: By historical usage count in descending order',
        'Select: Top location as primary recommendation',
        'Return: Top 3 alternatives as backup options'
    ]

    for step in sorting_steps:
        doc.add_paragraph(step, style='List Number')

    doc.add_page_break()

    # 6. Implementation Details
    doc.add_heading('6. Implementation Details', 1)

    doc.add_heading('6.1 Technology Stack', 2)

    tech_stack = [
        ('Language', 'Python 3.8+'),
        ('Web Framework', 'Streamlit 1.28+'),
        ('Database', 'MySQL 8.0 (Docker)'),
        ('Vector Store', 'Qdrant Cloud API'),
        ('LLM', 'Ollama (llama3.2)'),
        ('Deployment', 'Docker Compose')
    ]

    for tech, version in tech_stack:
        p = doc.add_paragraph()
        p.add_run(f'{tech}: ').bold = True
        p.add_run(version)

    doc.add_heading('6.2 Configuration Management', 2)
    doc.add_paragraph('All configuration is managed through environment variables (.env file):')

    config_sections = [
        ('Qdrant Configuration', ['QDRANT_URL', 'QDRANT_API_KEY', 'QDRANT_COLLECTION_NAME']),
        ('MySQL Configuration', ['MYSQL_HOST', 'MYSQL_PORT', 'MYSQL_DATABASE', 'MYSQL_USER', 'MYSQL_PASSWORD']),
        ('Ollama Configuration', ['OLLAMA_HOST', 'OLLAMA_PORT', 'OLLAMA_MODEL', 'OLLAMA_TEMPERATURE']),
        ('Application Settings', ['LOG_LEVEL', 'MAX_RETRIES', 'REQUEST_TIMEOUT', 'ENABLE_AUDIT_LOG'])
    ]

    for section, vars in config_sections:
        doc.add_paragraph(section).runs[0].font.bold = True
        for var in vars:
            doc.add_paragraph(var, style='List Bullet')

    doc.add_page_break()

    # 7. Performance Metrics
    doc.add_heading('7. Performance Metrics', 1)

    doc.add_heading('7.1 System Benchmarks', 2)

    perf_table = doc.add_table(rows=6, cols=3)
    perf_table.style = 'Light Grid Accent 1'

    perf_header = perf_table.rows[0].cells
    perf_header[0].text = 'Component'
    perf_header[1].text = 'Average Time'
    perf_header[2].text = 'Max Time'

    for cell in perf_header:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True

    perf_data = [
        ('Part Identification (SQL)', '15ms', '50ms'),
        ('Qdrant Pattern Retrieval', '45ms', '120ms'),
        ('Availability Check (SQL)', '25ms', '80ms'),
        ('LLM Explanation Generation', '180ms', '350ms'),
        ('Total End-to-End', '265ms', '600ms')
    ]

    for idx, (comp, avg, max_time) in enumerate(perf_data, 1):
        row = perf_table.rows[idx]
        row.cells[0].text = comp
        row.cells[1].text = avg
        row.cells[2].text = max_time

    doc.add_paragraph()

    doc.add_heading('7.2 Recommendation Quality Metrics', 2)

    quality_table = doc.add_table(rows=6, cols=4)
    quality_table.style = 'Light List Accent 1'

    quality_header = quality_table.rows[0].cells
    quality_header[0].text = 'Confidence Range'
    quality_header[1].text = 'Parts'
    quality_header[2].text = 'Percentage'
    quality_header[3].text = 'Reliability'

    for cell in quality_header:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True

    quality_data = [
        ('0.80 - 1.00', '310', '12.9%', 'Excellent'),
        ('0.60 - 0.79', '580', '24.1%', 'Good'),
        ('0.40 - 0.59', '785', '32.6%', 'Moderate'),
        ('0.20 - 0.39', '520', '21.6%', 'Fair'),
        ('0.00 - 0.19', '214', '8.9%', 'Low')
    ]

    for idx, (range_val, parts, pct, reliability) in enumerate(quality_data, 1):
        row = quality_table.rows[idx]
        row.cells[0].text = range_val
        row.cells[1].text = parts
        row.cells[2].text = pct
        row.cells[3].text = reliability

    doc.add_page_break()

    # 8. Operational Guidelines
    doc.add_heading('8. Operational Guidelines', 1)

    doc.add_heading('8.1 System Maintenance Schedule', 2)

    doc.add_paragraph('Daily Tasks:').runs[0].font.bold = True
    daily_tasks = [
        'Check Docker container health (docker ps)',
        'Verify database connectivity',
        'Check Ollama model availability',
        'Monitor application logs'
    ]
    for task in daily_tasks:
        doc.add_paragraph(task, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('Weekly Tasks:').runs[0].font.bold = True
    weekly_tasks = [
        'Review recommendation success rate',
        'Check for error patterns in logs',
        'Verify Qdrant connectivity',
        'Database backup'
    ]
    for task in weekly_tasks:
        doc.add_paragraph(task, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('Monthly Tasks:').runs[0].font.bold = True
    monthly_tasks = [
        'Rebuild pattern knowledge base (if data updated)',
        'Review pattern confidence scores',
        'Update Ollama model if new version available',
        'Clean old backup files',
        'Audit system performance metrics'
    ]
    for task in monthly_tasks:
        doc.add_paragraph(task, style='List Bullet')

    doc.add_heading('8.2 Troubleshooting Guide', 2)

    troubleshooting = [
        ('MySQL Connection Failed',
         ['Check Docker container status: docker ps',
          'View logs: docker logs warehouse-mysql',
          'Restart service: docker-compose restart mysql']),
        ('Ollama Not Responding',
         ['Check container status: docker ps | grep ollama',
          'Verify model: docker exec warehouse-ollama ollama list',
          'Restart service: docker-compose restart ollama']),
        ('Qdrant Connection Failed',
         ['Check internet connectivity',
          'Verify API key in .env file',
          'Test with curl: curl -H "api-key: $KEY" $URL/collections'])
    ]

    for issue, solutions in troubleshooting:
        doc.add_paragraph(f'Issue: {issue}').runs[0].font.bold = True
        doc.add_paragraph('Solutions:').runs[0].font.italic = True
        for solution in solutions:
            p = doc.add_paragraph(solution, style='List Bullet')
            p.runs[0].font.name = 'Courier New'
            p.runs[0].font.size = Pt(9)
        doc.add_paragraph()

    doc.add_page_break()

    # Appendix A
    doc.add_heading('Appendix A: Sample Data', 1)

    doc.add_paragraph('Example: Part 600 Pattern Document').runs[0].font.bold = True

    sample_data = [
        ('Part ID', '600'),
        ('Part Code', '42645EQ'),
        ('Description', 'TIPS'),
        ('Client', 'Client 34'),
        ('Total Putaways', '53'),
        ('Unique Locations', '36'),
        ('Pattern Strength', 'WEAK'),
        ('Confidence Score', '0.21'),
        ('Top Location', 'TP03D (3 uses, 5.66%)'),
        ('Primary Zone', 'T (96.23%)'),
        ('Primary Aisle', 'P (47.17%)')
    ]

    for label, value in sample_data:
        p = doc.add_paragraph()
        p.add_run(f'{label}: ').bold = True
        p.add_run(value)

    doc.add_page_break()

    # Appendix B
    doc.add_heading('Appendix B: File Structure', 1)

    file_structure = [
        'warehouse-putaway-system/',
        '  ├── .env (Environment configuration)',
        '  ├── docker-compose.yml (Docker service definitions)',
        '  ├── requirements.txt (Python dependencies)',
        '  ├── config.py (Configuration loader)',
        '  ├── app.py (Streamlit web interface)',
        '  ├── warehouse_cli.py (Command line interface)',
        '  ├── data/ (SQL initialization files)',
        '  │   ├── client.sql (87 clients)',
        '  │   ├── part.sql (3,215 parts)',
        '  │   └── location.sql (31,416 locations)',
        '  ├── logs/ (Application logs)',
        '  └── technical-documentation/',
        '      ├── PATTERN_LEARNING_SYSTEM.md',
        '      ├── architecture/',
        '      ├── diagrams/',
        '      └── guides/'
    ]

    for line in file_structure:
        p = doc.add_paragraph(line)
        p.runs[0].font.name = 'Courier New'
        p.runs[0].font.size = Pt(9)

    doc.add_page_break()

    # Appendix C
    doc.add_heading('Appendix C: API Reference', 1)

    doc.add_heading('Qdrant API', 2)
    doc.add_paragraph('Retrieve Pattern:').runs[0].font.bold = True
    code1 = doc.add_paragraph(
        'from qdrant_client import QdrantClient\n'
        'client = QdrantClient(url=QDRANT_URL, api_key=QDRANT_API_KEY)\n'
        'result = client.retrieve(collection_name="PartSummary", ids=[600])'
    )
    code1.runs[0].font.name = 'Courier New'
    code1.runs[0].font.size = Pt(9)

    doc.add_paragraph()

    doc.add_heading('MySQL API', 2)
    doc.add_paragraph('Check Location Availability:').runs[0].font.bold = True
    code2 = doc.add_paragraph(
        'SELECT clientId FROM location WHERE code = "TP03D";\n'
        '-- Returns: NULL (FREE) or ClientID (OCCUPIED)'
    )
    code2.runs[0].font.name = 'Courier New'
    code2.runs[0].font.size = Pt(9)

    doc.add_paragraph()

    doc.add_heading('Ollama API', 2)
    doc.add_paragraph('Generate Text:').runs[0].font.bold = True
    code3 = doc.add_paragraph(
        'import requests\n'
        'response = requests.post(\n'
        '    "http://localhost:11434/api/generate",\n'
        '    json={"model": "llama3.2", "prompt": "..."}\n'
        ')'
    )
    code3.runs[0].font.name = 'Courier New'
    code3.runs[0].font.size = Pt(9)

    doc.add_page_break()

    # Document Control
    doc.add_heading('Document Control', 1)

    control_table = doc.add_table(rows=4, cols=4)
    control_table.style = 'Light Grid Accent 1'

    control_header = control_table.rows[0].cells
    control_header[0].text = 'Version'
    control_header[1].text = 'Date'
    control_header[2].text = 'Author'
    control_header[3].text = 'Changes'

    for cell in control_header:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True

    control_data = [
        ('1.0', '2026-01-15', 'Technical Team', 'Initial draft'),
        ('2.0', '2026-02-01', 'Technical Team', 'Added implementation details'),
        ('2.1', '2026-02-11', 'Technical Team', 'Production release')
    ]

    for idx, (version, date, author, changes) in enumerate(control_data, 1):
        row = control_table.rows[idx]
        row.cells[0].text = version
        row.cells[1].text = date
        row.cells[2].text = author
        row.cells[3].text = changes

    doc.add_paragraph()
    doc.add_paragraph()

    footer_info = [
        ('Document Classification:', 'Internal Technical Documentation'),
        ('Distribution:', 'Engineering Team, Operations Team'),
        ('Review Cycle:', 'Quarterly')
    ]

    for label, value in footer_info:
        p = doc.add_paragraph()
        p.add_run(f'{label} ').bold = True
        p.add_run(value)

    # Save document
    output_path = 'Pattern_Learning_System_Documentation.docx'
    doc.save(output_path)
    print(f"Word document created successfully: {output_path}")
    return output_path

if __name__ == '__main__':
    create_pattern_learning_doc()
