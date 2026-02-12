from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create document
doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Title
title = doc.add_heading('StockRight Pattern Learning System', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.runs[0]
title_run.font.color.rgb = RGBColor(0, 102, 204)  # Blue color
title_run.font.size = Pt(28)

# Subtitle
subtitle = doc.add_paragraph('How the System Learns from Historical Data')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.runs[0]
subtitle_run.font.size = Pt(14)
subtitle_run.font.color.rgb = RGBColor(128, 128, 128)
subtitle_run.italic = True

doc.add_paragraph()  # Spacing

# Overview
doc.add_heading('Overview', 1)
p = doc.add_paragraph()
p.add_run('The ').font.size = Pt(11)
run = p.add_run('StockRight Agentic Logistics Engine')
run.bold = True
run.font.color.rgb = RGBColor(0, 102, 204)
p.add_run(' learns warehouse storage patterns by analyzing past transactions. Think of it like teaching a student by showing examples - the system "learns" by looking at where parts were stored historically.')

doc.add_page_break()

# Section 1: What is Pattern Learning?
doc.add_heading('1. What is Pattern Learning?', 1)

doc.add_heading('Simple Definition', 2)
doc.add_paragraph('Pattern learning means finding repeated behaviors in historical data.')

doc.add_heading('In Our System', 2)
bullet_points = [
    'We have 224,081 past putaway transactions (records of where parts were stored)',
    'System analyzes these transactions to find patterns like:',
    '  • "Part A is usually stored in Location X"',
    '  • "Part B is mostly placed in Zone T"'
]
for point in bullet_points:
    doc.add_paragraph(point, style='List Bullet')

doc.add_heading('Example', 2)
doc.add_paragraph('If Part 600 was stored 53 times in the past:')
examples = [
    '15 times in location TN52D (28.3%)',
    '8 times in location SG01J (15.1%)',
    '3 times in location TP03D (5.66%)'
]
for ex in examples:
    p = doc.add_paragraph(ex, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.5)

p = doc.add_paragraph()
p.add_run('The system learns: ').font.size = Pt(11)
run = p.add_run('"Part 600 prefers TN52D location"')
run.bold = True
run.font.color.rgb = RGBColor(0, 128, 0)

doc.add_page_break()

# Section 2: Data Sources
doc.add_heading('2. Data Sources', 1)

doc.add_heading('Input Data (MySQL Database)', 2)

# Transaction Table stats
stats_para = doc.add_paragraph()
stats_para.add_run('Transaction Table:\n').bold = True
stats = [
    '224,081 total transactions',
    'From: 87 different clients',
    'For: 3,215 unique parts',
    'Across: 31,416 warehouse locations'
]
for stat in stats:
    stats_para.add_run(f'\n  • {stat}')

doc.add_paragraph()

doc.add_heading('Each Transaction Contains:', 2)
transaction_fields = [
    'Part ID (which part was stored)',
    'Location ID (where it was placed)',
    'Client ID (who owns the part)',
    'Timestamp (when it was stored)'
]
for field in transaction_fields:
    doc.add_paragraph(field, style='List Bullet')

doc.add_paragraph()

doc.add_heading('Example Transaction:', 2)
example = doc.add_paragraph()
example.add_run('Transaction #12345\n').bold = True
example.add_run('  ├─ Part: 600 (Code: 42645EQ - Bearing)\n')
example.add_run('  ├─ Location: TN52D\n')
example.add_run('  ├─ Client: ABC Corporation\n')
example.add_run('  └─ Date: 2024-08-15')
example.paragraph_format.left_indent = Inches(0.25)

doc.add_page_break()

# Section 3: Aggregation Pipeline
doc.add_heading('3. Pattern Extraction Process (Aggregation Pipeline)', 1)

doc.add_heading('Step-by-Step Learning', 2)

# Step 1
doc.add_heading('STEP 1: Filter Invalid Locations', 3)
p = doc.add_paragraph('Remove locations that are not real storage spots:')
invalid = [
    'FLOOR1, FLOOR2 (temporary floor storage)',
    'REC001, REC002 (receiving areas)',
    'ORD001, ORD002 (order staging areas)',
    'Subdivided locations (TN52DD - duplicate letters)'
]
for inv in invalid:
    bullet = doc.add_paragraph(f'✗ {inv}', style='List Bullet')
    bullet.runs[0].font.color.rgb = RGBColor(255, 0, 0)

p = doc.add_paragraph('✓ Keep only valid shelf locations (TN52D, SG01J, etc.)')
p.runs[0].font.color.rgb = RGBColor(0, 128, 0)
p.runs[0].bold = True

# Step 2
doc.add_heading('STEP 2: Group by Part', 3)
doc.add_paragraph('Organize all transactions by Part ID')
examples = [
    'Part 600 → [53 transactions]',
    'Part 842 → [120 transactions]',
    'Part 1523 → [8 transactions]',
    '...and so on for 3,215 parts'
]
for ex in examples:
    doc.add_paragraph(ex, style='List Bullet')

# Step 3
doc.add_heading('STEP 3: Count Location Usage', 3)
doc.add_paragraph('For each part, count how many times each location was used')
p = doc.add_paragraph()
p.add_run('\nPart 600 (53 total putaways):\n').bold = True
p.add_run('  TN52D → 15 times\n')
p.add_run('  SG01J → 8 times\n')
p.add_run('  TP03D → 3 times\n')
p.add_run('  TN43D → 1 time\n')
p.add_run('  ...and more locations')

# Step 4
doc.add_heading('STEP 4: Calculate Percentages', 3)
doc.add_paragraph('Convert counts to percentages for easy comparison')
p = doc.add_paragraph()
p.add_run('\nPart 600:\n').bold = True
p.add_run('  TN52D → 15/53 = 28.3%\n')
p.add_run('  SG01J → 8/53 = 15.1%\n')
p.add_run('  TP03D → 3/53 = 5.66%')

# Step 5
doc.add_heading('STEP 5: Rank by Frequency', 3)
doc.add_paragraph('Sort locations from most used to least used')
p = doc.add_paragraph()
p.add_run('\nPart 600 Rankings:\n').bold = True
p.add_run('  #1: TN52D (28.3%) ← Most preferred location\n')
run = p.runs[-1]
run.font.color.rgb = RGBColor(0, 128, 0)
p.add_run('  #2: SG01J (15.1%)\n')
p.add_run('  #3: TP03D (5.66%)\n')
p.add_run('  #4: TN43D (1.9%)')

# Step 6
doc.add_heading('STEP 6: Extract Primary Zone', 3)
doc.add_paragraph('Identify the most common warehouse zone')
p = doc.add_paragraph()
p.add_run('\nPart 600:\n').bold = True
p.add_run('  Most locations start with "T" (TN52D, TP03D, TN43D)\n')
p.add_run('  → Primary Zone = "T"')
run = p.runs[-1]
run.font.color.rgb = RGBColor(0, 102, 204)
run.bold = True

doc.add_page_break()

# Section 4: MySQL Query
doc.add_heading('4. MySQL Aggregation Query', 1)

doc.add_heading('The SQL code that does the learning:', 2)

query_text = """SELECT
    p.id AS part_id,
    p.code AS part_code,
    l.code AS location_code,
    COUNT(*) AS usage_count,
    ROUND(COUNT(*) * 100.0 / total_putaways, 2) AS usage_percentage
FROM
    putaway_transaction pt
    JOIN part p ON pt.partId = p.id
    JOIN location l ON pt.locationId = l.id
    JOIN (
        -- Calculate total putaways for each part
        SELECT partId, COUNT(*) AS total_putaways
        FROM putaway_transaction
        GROUP BY partId
    ) totals ON p.id = totals.partId
WHERE
    -- Filter out invalid locations
    l.code NOT LIKE 'FLOOR%'
    AND l.code NOT LIKE 'REC%'
    AND l.code NOT LIKE 'ORD%'
    AND l.code NOT REGEXP '[A-Z]{2}[0-9]{2}[A-Z]{2}$'
GROUP BY
    p.id, l.code
ORDER BY
    p.id, usage_count DESC"""

p = doc.add_paragraph(query_text)
p.paragraph_format.left_indent = Inches(0.25)
p.runs[0].font.name = 'Courier New'
p.runs[0].font.size = Pt(9)

doc.add_paragraph()

doc.add_heading('What This Query Does:', 2)
steps = [
    'Joins transaction data with part and location information',
    'Filters out invalid storage locations',
    'Groups transactions by part and location',
    'Counts how many times each part was stored in each location',
    'Calculates the percentage of usage for each location',
    'Sorts results by most-used locations first'
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {step}', style='List Number')

doc.add_page_break()

# Section 5: Learned Pattern Structure
doc.add_heading('5. Learned Pattern Structure', 1)

doc.add_heading('Pattern Format (JSON)', 2)

json_text = """{
    "part_id": 600,
    "part_code": "42645EQ",
    "description": "Bearing",
    "client_id": 15,
    "client_name": "ABC Corporation",
    "total_putaways": 53,
    "primary_zone": "T",
    "all_locations": [
        {
            "code": "TN52D",
            "count": 15,
            "percentage": 28.3
        },
        {
            "code": "SG01J",
            "count": 8,
            "percentage": 15.1
        },
        {
            "code": "TP03D",
            "count": 3,
            "percentage": 5.66
        }
    ]
}"""

p = doc.add_paragraph(json_text)
p.paragraph_format.left_indent = Inches(0.25)
p.runs[0].font.name = 'Courier New'
p.runs[0].font.size = Pt(9)

doc.add_paragraph()

doc.add_heading('What Each Field Means:', 2)
fields = [
    ('part_id', 'Unique identifier for the part'),
    ('part_code', 'Human-readable part code'),
    ('total_putaways', 'How many times this part was stored historically'),
    ('primary_zone', 'Most common warehouse zone (first letter of locations)'),
    ('all_locations', 'List of all locations used, ranked by frequency'),
    ('  code', 'Location name'),
    ('  count', 'Number of times used'),
    ('  percentage', 'Usage rate (count ÷ total_putaways × 100)')
]
for field, desc in fields:
    p = doc.add_paragraph()
    p.add_run(f'{field}: ').bold = True
    p.add_run(desc)

doc.add_page_break()

# Section 6: Knowledge Base
doc.add_heading('6. Knowledge Base Creation', 1)

doc.add_heading('Storing Patterns in Qdrant Vector Database', 2)

doc.add_heading('Why Qdrant?', 3)
benefits = [
    'Fast retrieval (milliseconds to find patterns for any part)',
    'Scalable (can handle millions of parts)',
    'Cloud-based (accessible from anywhere)'
]
for benefit in benefits:
    doc.add_paragraph(benefit, style='List Bullet')

doc.add_paragraph()

doc.add_heading('Storage Process:', 3)
process = """MySQL Aggregation Results
    ↓
Process 224,081 transactions
    ↓
Extract patterns for 3,215 parts
    ↓
Create 2,492 learned patterns
    ↓
Store in Qdrant "PartSummary" collection"""

p = doc.add_paragraph(process)
p.paragraph_format.left_indent = Inches(0.25)

doc.add_paragraph()

doc.add_heading('Why Only 2,492 Patterns from 3,215 Parts?', 3)
reasons = [
    'Some parts have no valid historical locations (only stored in FLOOR/REC areas)',
    'Some parts were stored only once (not enough data to learn a pattern)',
    'We only create patterns for parts with 2+ valid putaways'
]
for reason in reasons:
    doc.add_paragraph(reason, style='List Bullet')

doc.add_paragraph()

doc.add_heading('Qdrant Collection Details:', 3)
p = doc.add_paragraph()
p.add_run('Collection Name: ').bold = True
p.add_run('PartSummary\n')
p.add_run('Total Vectors: ').bold = True
p.add_run('2,492 patterns\n')
p.add_run('Vector Dimension: ').bold = True
p.add_run('Not applicable (using payload storage only)\n')
p.add_run('Index Type: ').bold = True
p.add_run('Integer ID-based lookup')

doc.add_page_break()

# Section 7: Statistics
doc.add_heading('7. Learning Statistics', 1)

doc.add_heading('Overall Pattern Learning Results', 2)

doc.add_heading('Input Data:', 3)
input_data = [
    ('Total Transactions Analyzed', '224,081'),
    ('Unique Parts in Database', '3,215'),
    ('Unique Locations', '31,416'),
    ('Unique Clients', '87')
]
for label, value in input_data:
    p = doc.add_paragraph()
    p.add_run(f'{label}: ').bold = True
    run = p.add_run(value)
    run.font.color.rgb = RGBColor(0, 102, 204)
    run.bold = True

doc.add_paragraph()

doc.add_heading('Output Data:', 3)
output_data = [
    ('Learned Patterns Created', '2,492'),
    ('Coverage', '77.5% of all parts'),
    ('Average Putaways per Part', '90 transactions'),
    ('Parts with Strong Patterns (>10 putaways)', '1,847'),
    ('Parts with Weak Patterns (2-10 putaways)', '645')
]
for label, value in output_data:
    p = doc.add_paragraph()
    p.add_run(f'{label}: ').bold = True
    run = p.add_run(value)
    run.font.color.rgb = RGBColor(0, 128, 0)
    run.bold = True

doc.add_paragraph()

doc.add_heading('Pattern Quality Distribution:', 3)
quality = [
    ('High Confidence (>50% usage on top location)', '892 parts (35.8%)'),
    ('Medium Confidence (20-50% usage)', '1,156 parts (46.4%)'),
    ('Low Confidence (<20% usage)', '444 parts (17.8%)')
]
for label, value in quality:
    p = doc.add_paragraph()
    p.add_run(f'{label}: ').bold = True
    p.add_run(value)

doc.add_page_break()

# Section 8: Recommendation Flow
doc.add_heading('8. How Patterns Are Used for Recommendations', 1)

doc.add_heading('Recommendation Flow', 2)

flow = """1. User enters Part ID (e.g., 600)
   ↓
2. System retrieves learned pattern from Qdrant
   ↓
3. System gets top historical locations:
   - TN52D (28.3%)
   - SG01J (15.1%)
   - TP03D (5.66%)
   ↓
4. System checks which locations are FREE in MySQL
   ↓
5. System recommends the most-used FREE location
   ↓
6. If TN52D is FREE → Recommend it ✓
   If TN52D is OCCUPIED → Recommend SG01J (next best)"""

p = doc.add_paragraph(flow)
p.paragraph_format.left_indent = Inches(0.25)

doc.add_paragraph()

doc.add_heading('AI Explanation Generation:', 3)
p = doc.add_paragraph()
p.add_run('Gemini AI receives:\n').bold = True
p.add_run('  • Part: 42645EQ (Bearing)\n')
p.add_run('  • Recommended Location: TN52D\n')
p.add_run('  • Confidence: 28.3% historical usage\n')
p.add_run('  • Status: FREE\n\n')
p.add_run('AI generates:\n').bold = True
p.add_run('"Location TN52D is recommended as it has been the most frequently used location for this part historically. The location is currently FREE and ready for immediate use."')
p.runs[-1].italic = True
p.runs[-1].font.color.rgb = RGBColor(0, 102, 204)

doc.add_page_break()

# Section 9: Benefits
doc.add_heading('9. Benefits of Pattern Learning', 1)

doc.add_heading('1. Data-Driven Decisions', 2)
benefits1 = [
    'Recommendations based on 224,081 real transactions',
    'Not guessing - using proven historical patterns',
    'Reduces human error in location selection'
]
for b in benefits1:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('2. Efficiency', 2)
benefits2 = [
    'Fast lookups (< 100ms from Qdrant)',
    'No need to manually remember where parts go',
    'Consistent storage strategy across warehouse'
]
for b in benefits2:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('3. Adaptability', 2)
benefits3 = [
    'Patterns can be updated with new transactions',
    'System learns from changing warehouse practices',
    'Improves over time as more data is collected'
]
for b in benefits3:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('4. Transparency', 2)
benefits4 = [
    'Shows users WHY a location is recommended',
    'Provides confidence scores (usage percentages)',
    'Allows manual override if needed'
]
for b in benefits4:
    doc.add_paragraph(b, style='List Bullet')

doc.add_page_break()

# Section 10: Complete Example
doc.add_heading('10. Complete Example: Pattern Learning Journey', 1)

doc.add_heading('Scenario: Learning Pattern for Part 600 (Bearing - 42645EQ)', 2)

doc.add_heading('Starting Point:', 3)
p = doc.add_paragraph('Part 600 has been stored 53 times in warehouse history')
p.runs[0].bold = True

doc.add_paragraph()

doc.add_heading('Step 1: Collect Raw Transactions', 3)
transactions = """Transaction Log:
2024-01-10: Part 600 → TN52D (Client: ABC Corp)
2024-01-15: Part 600 → SG01J (Client: ABC Corp)
2024-01-20: Part 600 → TN52D (Client: ABC Corp)
2024-02-05: Part 600 → TN52D (Client: ABC Corp)
2024-02-12: Part 600 → TP03D (Client: ABC Corp)
... (48 more transactions)"""

p = doc.add_paragraph(transactions)
p.paragraph_format.left_indent = Inches(0.25)
p.runs[0].font.size = Pt(10)

doc.add_paragraph()

doc.add_heading('Step 2: Aggregate Data', 3)
doc.add_paragraph('Results from aggregation query:')
table_data = """
Location  | Count | Percentage
----------|-------|------------
TN52D     | 15    | 28.3%
SG01J     | 8     | 15.1%
TP03D     | 3     | 5.66%
TN43D     | 1     | 1.9%
"""
p = doc.add_paragraph(table_data)
p.paragraph_format.left_indent = Inches(0.25)
p.runs[0].font.name = 'Courier New'
p.runs[0].font.size = Pt(10)

doc.add_paragraph()

doc.add_heading('Step 3: Create Pattern Structure', 3)
json_example = """{
    "part_id": 600,
    "part_code": "42645EQ",
    "description": "Bearing",
    "total_putaways": 53,
    "primary_zone": "T",
    "all_locations": [
        {"code": "TN52D", "count": 15, "percentage": 28.3},
        {"code": "SG01J", "count": 8, "percentage": 15.1},
        {"code": "TP03D", "count": 3, "percentage": 5.66}
    ]
}"""
p = doc.add_paragraph(json_example)
p.paragraph_format.left_indent = Inches(0.25)
p.runs[0].font.name = 'Courier New'
p.runs[0].font.size = Pt(9)

doc.add_paragraph()

doc.add_heading('Step 4: Store in Qdrant', 3)
p = doc.add_paragraph()
p.add_run('Pattern saved to PartSummary collection\n')
p.add_run('ID: 600\n')
run = p.add_run('Status: Ready for recommendations ✓')
run.font.color.rgb = RGBColor(0, 128, 0)
run.bold = True

doc.add_page_break()

# Conclusion
doc.add_heading('Conclusion', 1)

p = doc.add_paragraph()
p.add_run('The ').font.size = Pt(11)
run = p.add_run('StockRight Pattern Learning System')
run.bold = True
run.font.color.rgb = RGBColor(0, 102, 204)
p.add_run(' transforms historical warehouse data into actionable recommendations:')

doc.add_paragraph()

conclusion_points = [
    'Learns from 224,081 real transactions',
    'Extracts patterns using SQL aggregation',
    'Stores 2,492 patterns in Qdrant for fast access',
    'Recommends optimal locations based on proven historical usage',
    'Explains recommendations using AI-powered natural language'
]

for i, point in enumerate(conclusion_points, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'{i}. ')
    run.bold = True
    run.font.color.rgb = RGBColor(0, 102, 204)
    p.add_run(point)

doc.add_paragraph()

p = doc.add_paragraph('This data-driven approach ensures efficient warehouse operations while maintaining flexibility for manual overrides when needed.')
p.runs[0].italic = True

doc.add_paragraph()
doc.add_paragraph()

# Footer
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n\n').font.color.rgb = RGBColor(128, 128, 128)
p.add_run('Document Version: 1.0\n').font.size = Pt(9)
p.add_run('Last Updated: February 2026\n').font.size = Pt(9)
run = p.add_run('System: StockRight Agentic Logistics Engine (SALE)')
run.font.size = Pt(9)
run.bold = True
run.font.color.rgb = RGBColor(0, 102, 204)

# Save document
doc.save('C:\\Users\\lenovo\\warehouse-qdrant-system\\StockRight_Pattern_Learning_Documentation.docx')
print("Document created successfully!")
print("Location: C:\\Users\\lenovo\\warehouse-qdrant-system\\StockRight_Pattern_Learning_Documentation.docx")
