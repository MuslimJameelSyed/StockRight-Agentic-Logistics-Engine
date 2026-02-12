from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_page_number(doc):
    """Add page numbers to footer"""
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = footer_para.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    run._r.append(instrText)

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar2)

# Create document
doc = Document()

# Set margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# =============================================================================
# COVER PAGE
# =============================================================================

# Logo/Title area
title = doc.add_heading('StockRight', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.runs[0]
title_run.font.color.rgb = RGBColor(0, 102, 204)
title_run.font.size = Pt(36)
title_run.font.name = 'Arial'

# Main Title
subtitle = doc.add_paragraph('Pattern Learning System')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.runs[0]
subtitle_run.font.size = Pt(24)
subtitle_run.font.color.rgb = RGBColor(0, 102, 204)
subtitle_run.bold = True

# Subtitle
tagline = doc.add_paragraph('Knowledge Base & Aggregation Pipeline Documentation')
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
tagline_run = tagline.runs[0]
tagline_run.font.size = Pt(14)
tagline_run.font.color.rgb = RGBColor(128, 128, 128)
tagline_run.italic = True

doc.add_paragraph('\n\n\n')

# Document info box
info_para = doc.add_paragraph()
info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
info_para.add_run('━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n\n').font.color.rgb = RGBColor(200, 200, 200)
info_run = info_para.add_run('System: StockRight Agentic Logistics Engine (SALE)\n')
info_run.font.size = Pt(11)
info_run = info_para.add_run('Version: 1.0\n')
info_run.font.size = Pt(11)
info_run = info_para.add_run('Date: February 2026\n')
info_run.font.size = Pt(11)
info_run = info_para.add_run('Knowledge Base: 224,081 Transactions | 2,492 Patterns\n\n')
info_run.font.size = Pt(11)
info_run.bold = True
info_run.font.color.rgb = RGBColor(0, 102, 204)
info_para.add_run('━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━').font.color.rgb = RGBColor(200, 200, 200)

doc.add_page_break()

# =============================================================================
# TABLE OF CONTENTS
# =============================================================================

toc_heading = doc.add_heading('Table of Contents', 1)
toc_heading.runs[0].font.color.rgb = RGBColor(0, 102, 204)

toc_items = [
    ('1. Executive Summary', 3),
    ('2. System Overview', 3),
    ('3. Data Sources & Statistics', 3),
    ('4. Pattern Learning Process (6-Step Aggregation)', 4),
    ('5. MySQL Aggregation Query', 5),
    ('6. Knowledge Base Structure', 6),
    ('7. What is Stored vs. Real-Time Queries', 7),
    ('8. Hybrid Architecture (Offline + Online)', 8),
    ('9. Pattern Quality & Statistics', 9),
    ('10. Recommendation Flow', 10),
    ('11. Complete Example: Part 600 Journey', 11),
    ('12. Benefits & Advantages', 12),
    ('13. Update Schedule & Maintenance', 13)
]

for item, page in toc_items:
    p = doc.add_paragraph()
    p.add_run(item).font.size = Pt(11)
    p.add_run(' ' + '.' * 50 + ' ').font.color.rgb = RGBColor(200, 200, 200)
    p.add_run(str(page)).font.size = Pt(11)
    p.paragraph_format.left_indent = Inches(0.25)

doc.add_page_break()

# =============================================================================
# SECTION 1: EXECUTIVE SUMMARY
# =============================================================================

doc.add_heading('1. Executive Summary', 1)

p = doc.add_paragraph()
p.add_run('The ').font.size = Pt(11)
run = p.add_run('StockRight Pattern Learning System')
run.bold = True
run.font.color.rgb = RGBColor(0, 102, 204)
p.add_run(' is an AI-powered warehouse recommendation engine that learns from historical data to suggest optimal storage locations. The system combines offline pattern analysis with real-time availability checks to provide intelligent, data-driven recommendations.')

doc.add_paragraph()

doc.add_heading('Key Highlights', 2)

highlights = [
    ('224,081', 'historical transactions analyzed'),
    ('2,492', 'learned patterns created'),
    ('3,215', 'unique parts tracked'),
    ('31,416', 'warehouse locations'),
    ('77.5%', 'pattern coverage'),
    ('<100ms', 'recommendation speed')
]

for number, description in highlights:
    p = doc.add_paragraph()
    run = p.add_run(number + ' ')
    run.bold = True
    run.font.color.rgb = RGBColor(0, 128, 0)
    run.font.size = Pt(14)
    p.add_run('- ' + description)

doc.add_page_break()

# =============================================================================
# SECTION 2: SYSTEM OVERVIEW
# =============================================================================

doc.add_heading('2. System Overview', 1)

doc.add_heading('What is Pattern Learning?', 2)

p = doc.add_paragraph()
run = p.add_run('Pattern learning')
run.bold = True
p.add_run(' means finding repeated behaviors in historical data. Think of it like teaching a student by showing examples - the system "learns" by analyzing where parts were stored in the past.')

doc.add_paragraph()

doc.add_heading('Simple Example', 3)

example_box = doc.add_paragraph()
example_box.add_run('If Part 600 was stored 53 times historically:\n\n').bold = True
example_box.add_run('  • 15 times in location TN52D (28.3%)\n')
example_box.add_run('  • 8 times in location SG01J (15.1%)\n')
example_box.add_run('  • 3 times in location TP03D (5.66%)\n\n')
run = example_box.add_run('The system learns: "Part 600 prefers TN52D location"')
run.italic = True
run.font.color.rgb = RGBColor(0, 128, 0)

doc.add_paragraph()

doc.add_heading('Three-Layer Architecture', 2)

arch_text = """┌─────────────────────────────────────────┐
│   LAYER 1: DATA COLLECTION              │
│   MySQL Database (224,081 transactions) │
└────────────────┬────────────────────────┘
                 │
                 ▼
┌─────────────────────────────────────────┐
│   LAYER 2: PATTERN LEARNING             │
│   Aggregation Pipeline (6 steps)        │
│   → Filter → Group → Count → Calculate  │
└────────────────┬────────────────────────┘
                 │
                 ▼
┌─────────────────────────────────────────┐
│   LAYER 3: KNOWLEDGE BASE               │
│   Qdrant Vector DB (2,492 patterns)     │
│   + Real-Time MySQL Queries             │
└─────────────────────────────────────────┘"""

p = doc.add_paragraph(arch_text)
p.paragraph_format.left_indent = Inches(0.25)
p.runs[0].font.name = 'Courier New'
p.runs[0].font.size = Pt(9)

doc.add_page_break()

# =============================================================================
# SECTION 3: DATA SOURCES
# =============================================================================

doc.add_heading('3. Data Sources & Statistics', 1)

doc.add_heading('Input Data (MySQL Database)', 2)

# Create a nice stats box
stats_para = doc.add_paragraph()
stats_para.add_run('Transaction Dataset Overview:\n\n').bold = True
stats_para.runs[0].font.color.rgb = RGBColor(0, 102, 204)

stats = [
    ('Total Transactions', '224,081'),
    ('Unique Parts', '3,215'),
    ('Warehouse Locations', '31,416'),
    ('Clients', '87'),
    ('Date Range', '2023 - 2026')
]

for label, value in stats:
    stats_para.add_run(f'  • {label}: ')
    run = stats_para.add_run(value + '\n')
    run.bold = True
    run.font.color.rgb = RGBColor(0, 128, 0)

doc.add_paragraph()

doc.add_heading('Each Transaction Contains', 2)

fields = [
    ('Part ID', 'Which part was stored (e.g., 600)'),
    ('Location ID', 'Where it was placed (e.g., TN52D)'),
    ('Client ID', 'Who owns the part (e.g., Client 34)'),
    ('Timestamp', 'When it was stored (e.g., 2024-08-15)')
]

for field, desc in fields:
    p = doc.add_paragraph()
    run = p.add_run(f'{field}: ')
    run.bold = True
    run.font.color.rgb = RGBColor(0, 102, 204)
    p.add_run(desc)

doc.add_paragraph()

doc.add_heading('Example Transaction Record', 3)

transaction = """Transaction #12345
├─ Part: 600 (Code: 42645EQ - Bearing)
├─ Location: TN52D
├─ Client: ABC Corporation
└─ Date: 2024-08-15 14:30:00"""

p = doc.add_paragraph(transaction)
p.paragraph_format.left_indent = Inches(0.25)
p.runs[0].font.name = 'Courier New'
p.runs[0].font.size = Pt(10)

doc.add_page_break()

# =============================================================================
# SECTION 4: AGGREGATION PIPELINE (6 STEPS)
# =============================================================================

doc.add_heading('4. Pattern Learning Process (6-Step Aggregation)', 1)

intro = doc.add_paragraph()
intro.add_run('The aggregation pipeline is the ').font.size = Pt(11)
run = intro.add_run('heart of the learning process')
run.bold = True
run.font.color.rgb = RGBColor(0, 102, 204)
intro.add_run('. It transforms raw transaction data into meaningful patterns through six systematic steps.')

doc.add_paragraph()

# STEP 1
doc.add_heading('STEP 1: Filter Invalid Locations', 2)

p = doc.add_paragraph()
p.add_run('Remove locations that are NOT permanent storage:\n\n').bold = True

invalid_locations = [
    ('FLOOR1, FLOOR2', 'Temporary floor storage'),
    ('REC001, REC002', 'Receiving dock areas'),
    ('ORD001, ORD002', 'Order staging zones'),
    ('TN52DD, SG01AA', 'Subdivided locations (duplicate letters)')
]

for location, reason in invalid_locations:
    bullet = doc.add_paragraph(f'  X  {location} - {reason}', style='List Bullet')
    bullet.runs[0].font.color.rgb = RGBColor(255, 0, 0)

p = doc.add_paragraph()
run = p.add_run('✓ Keep only valid shelf locations: TN52D, SG01J, TP03D, etc.')
run.font.color.rgb = RGBColor(0, 128, 0)
run.bold = True

doc.add_paragraph()

# STEP 2
doc.add_heading('STEP 2: Group by Part ID', 2)

p = doc.add_paragraph('Organize all transactions by Part ID to analyze each part separately:\n')

grouping = """Part 600  →  [53 transactions]
Part 842  →  [120 transactions]
Part 1523 →  [8 transactions]
Part 2045 →  [95 transactions]
...
(Total: 3,215 unique parts)"""

p2 = doc.add_paragraph(grouping)
p2.paragraph_format.left_indent = Inches(0.5)
p2.runs[0].font.name = 'Courier New'
p2.runs[0].font.size = Pt(10)

doc.add_paragraph()

# STEP 3
doc.add_heading('STEP 3: Count Location Usage', 2)

p = doc.add_paragraph('For each part, count how many times each location was used:\n\n')

counting = """Part 600 (53 total putaways):
  • TN52D → Used 15 times
  • SG01J → Used 8 times
  • TP03D → Used 3 times
  • TN43D → Used 1 time
  • (+ 32 more locations with smaller counts)"""

p2 = doc.add_paragraph(counting)
p2.paragraph_format.left_indent = Inches(0.25)

doc.add_paragraph()

# STEP 4
doc.add_heading('STEP 4: Calculate Percentages', 2)

p = doc.add_paragraph('Convert raw counts to percentages for easier comparison:\n\n')

percentages = """Part 600 Calculations:
  • TN52D: 15 / 53 = 28.3% ← Highest usage
  • SG01J: 8 / 53 = 15.1%
  • TP03D: 3 / 53 = 5.66%
  • TN43D: 1 / 53 = 1.9%"""

p2 = doc.add_paragraph(percentages)
p2.paragraph_format.left_indent = Inches(0.25)
p2.runs[0].font.name = 'Courier New'
p2.runs[0].font.size = Pt(10)

doc.add_paragraph()

# STEP 5
doc.add_heading('STEP 5: Rank by Frequency', 2)

p = doc.add_paragraph('Sort locations from most used to least used:\n\n')

ranking = """Part 600 Rankings:
  #1: TN52D (28.3%) ← Primary recommendation
  #2: SG01J (15.1%) ← Alternative 1
  #3: TP03D (5.66%) ← Alternative 2
  #4: TN43D (1.9%)
  ..."""

p2 = doc.add_paragraph(ranking)
p2.paragraph_format.left_indent = Inches(0.25)

doc.add_paragraph()

# STEP 6
doc.add_heading('STEP 6: Extract Primary Zone', 2)

p = doc.add_paragraph('Identify the most common warehouse zone (first letter of location codes):\n\n')

zone_analysis = """Part 600 Zone Analysis:
  • TN52D → Zone "T"
  • TP03D → Zone "T"
  • TN43D → Zone "T"
  • SG01J → Zone "S"

Conclusion: Primary Zone = "T" (used 96.2% of the time)"""

p2 = doc.add_paragraph(zone_analysis)
p2.paragraph_format.left_indent = Inches(0.25)

run = doc.add_paragraph().add_run('Result: Part 600 strongly prefers Zone T locations')
run.italic = True
run.font.color.rgb = RGBColor(0, 102, 204)

doc.add_page_break()

# =============================================================================
# SECTION 5: MYSQL QUERY
# =============================================================================

doc.add_heading('5. MySQL Aggregation Query', 1)

doc.add_paragraph('This SQL query performs the entire aggregation pipeline:')
doc.add_paragraph()

query = """SELECT
    p.id AS part_id,
    p.code AS part_code,
    l.code AS location_code,
    COUNT(*) AS usage_count,
    ROUND(COUNT(*) * 100.0 / total_putaways, 2) AS usage_percentage
FROM
    putaway_transaction pt
    JOIN part p ON pt.partId = p.id
    JOIN location l ON pt.locationId = l.id
    JOIN (
        -- Subquery: Calculate total putaways for each part
        SELECT partId, COUNT(*) AS total_putaways
        FROM putaway_transaction
        GROUP BY partId
    ) totals ON p.id = totals.partId
WHERE
    -- Filter out invalid locations
    l.code NOT LIKE 'FLOOR%'
    AND l.code NOT LIKE 'REC%'
    AND l.code NOT LIKE 'ORD%'
    AND l.code NOT REGEXP '[A-Z]{2}[0-9]{2}[A-Z]{2}$'  -- No subdivisions
GROUP BY
    p.id, l.code
ORDER BY
    p.id, usage_count DESC"""

p = doc.add_paragraph(query)
p.paragraph_format.left_indent = Inches(0.25)
p.runs[0].font.name = 'Courier New'
p.runs[0].font.size = Pt(9)

doc.add_paragraph()

doc.add_heading('Query Explanation (Line by Line)', 2)

explanations = [
    ('Lines 1-5', 'SELECT clause: Choose which columns to return'),
    ('Lines 7-9', 'JOIN tables: Connect transactions with parts and locations'),
    ('Lines 10-13', 'Subquery: Calculate total putaways per part for percentage calculation'),
    ('Lines 15-18', 'WHERE clause: Filter out invalid storage locations'),
    ('Line 19', 'GROUP BY: Organize results by part and location'),
    ('Line 20', 'ORDER BY: Sort by part ID and usage count (highest first)')
]

for line_ref, explanation in explanations:
    p = doc.add_paragraph()
    run = p.add_run(f'{line_ref}: ')
    run.bold = True
    run.font.color.rgb = RGBColor(0, 102, 204)
    p.add_run(explanation)

doc.add_page_break()

# =============================================================================
# SECTION 6: KNOWLEDGE BASE STRUCTURE
# =============================================================================

doc.add_heading('6. Knowledge Base Structure', 1)

doc.add_heading('Pattern Format (JSON)', 2)

doc.add_paragraph('Each learned pattern is stored in this structured format:')
doc.add_paragraph()

json_pattern = """{
  "part_id": 600,
  "part_code": "42645EQ",
  "description": "Bearing",
  "client_id": 34,
  "client_name": "ABC Corporation",
  "total_putaways": 53,
  "primary_zone": "T",
  "zone_distribution": {
    "T": 96.23,
    "S": 3.77
  },
  "all_locations": [
    {
      "code": "TN52D",
      "zone": "T",
      "aisle": "N",
      "column": "52",
      "row": "D",
      "count": 15,
      "percentage": 28.3,
      "first_used": "2023-09-01",
      "last_used": "2024-08-15"
    },
    {
      "code": "SG01J",
      "count": 8,
      "percentage": 15.1
    }
  ],
  "pattern_strength": "MEDIUM",
  "confidence_score": 0.72
}"""

p = doc.add_paragraph(json_pattern)
p.paragraph_format.left_indent = Inches(0.25)
p.runs[0].font.name = 'Courier New'
p.runs[0].font.size = Pt(8.5)

doc.add_paragraph()

doc.add_heading('Field Descriptions', 2)

fields = [
    ('part_id', 'Unique identifier (used for fast Qdrant lookup)'),
    ('total_putaways', 'How many times this part was stored historically'),
    ('primary_zone', 'Most common warehouse zone (e.g., "T")'),
    ('all_locations', 'Ranked list of all historical locations'),
    ('count', 'Number of times a location was used'),
    ('percentage', 'Usage rate (count ÷ total_putaways × 100)'),
    ('pattern_strength', 'STRONG/MEDIUM/WEAK based on concentration'),
    ('confidence_score', 'Statistical confidence (0-1 scale)')
]

for field, desc in fields:
    p = doc.add_paragraph()
    run = p.add_run(f'{field}: ')
    run.bold = True
    run.font.color.rgb = RGBColor(0, 102, 204)
    p.add_run(desc)

doc.add_page_break()

# =============================================================================
# SECTION 7: STORED VS REAL-TIME
# =============================================================================

doc.add_heading('7. What is Stored vs. Real-Time Queries', 1)

intro_hybrid = doc.add_paragraph()
intro_hybrid.add_run('The system uses a ').font.size = Pt(11)
run = intro_hybrid.add_run('hybrid approach')
run.bold = True
run.font.color.rgb = RGBColor(0, 102, 204)
intro_hybrid.add_run(': some data is pre-computed and stored (fast), while critical data is queried in real-time (accurate).')

doc.add_paragraph()

doc.add_heading('✓ Stored in Knowledge Base (Qdrant)', 2)

stored_items = [
    'Historical putaway patterns (all dates, frequencies, percentages)',
    'Zone and aisle distributions per part',
    'Client preferences and tendencies',
    'Pattern strength analysis (STRONG/MEDIUM/WEAK)',
    'Confidence scores and statistical metrics',
    'Inventory snapshot (from knowledge base creation date)'
]

for item in stored_items:
    p = doc.add_paragraph(item, style='List Bullet')
    p.runs[0].font.color.rgb = RGBColor(0, 128, 0)

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('Why store these? ')
run.bold = True
p.add_run('They change slowly (weekly/monthly) and are expensive to compute. Pre-computing saves time during recommendations.')

doc.add_paragraph()

doc.add_heading('✗ NOT Stored - Queried Real-Time (MySQL)', 2)

realtime_items = [
    'Location availability (FREE/OCCUPIED/YOUR_LOCATION)',
    'Current inventory quantities',
    'Which client currently occupies each location',
    'Real-time space utilization',
    'Recent transactions (today, this week)'
]

for item in realtime_items:
    p = doc.add_paragraph(item, style='List Bullet')
    p.runs[0].font.color.rgb = RGBColor(255, 0, 0)

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('Why query these? ')
run.bold = True
p.add_run('They change constantly (hourly/daily). Storing them would cause stale data and conflicts. Real-time queries ensure accuracy.')

doc.add_page_break()

# =============================================================================
# SECTION 8: HYBRID ARCHITECTURE
# =============================================================================

doc.add_heading('8. Hybrid Architecture (Offline + Online)', 1)

doc.add_heading('How Recommendations Work', 2)

flow_diagram = """┌──────────────────────────────────────────────────┐
│  STEP 1: Load from Knowledge Base (Qdrant)      │
│  ────────────────────────────────────────────    │
│  • Part 600 historical patterns                  │
│  • Top locations: TN52D (28.3%), SG01J (15.1%)  │
│  • Primary zone: T                               │
│  • Client preferences                            │
│  Speed: < 10ms ⚡                                │
└───────────────────┬──────────────────────────────┘
                    │
                    ▼
┌──────────────────────────────────────────────────┐
│  STEP 2: Query Real-Time Data (MySQL)           │
│  ────────────────────────────────────────────    │
│  Query 1: Current inventory for Part 600         │
│    SELECT location, quantity                     │
│    FROM inventory WHERE partId = 600             │
│    Result: TN52D (600k units), TP55C (600k)     │
│                                                  │
│  Query 2: Check TN52D availability               │
│    SELECT clientId FROM location                 │
│    WHERE code = 'TN52D'                          │
│    Result: clientId = 34                         │
│                                                  │
│  Speed: ~50ms 🔄                                 │
└───────────────────┬──────────────────────────────┘
                    │
                    ▼
┌──────────────────────────────────────────────────┐
│  STEP 3: Combine & Validate                     │
│  ────────────────────────────────────────────    │
│  • Historical says: TN52D is best (28.3%)        │
│  • Real-time says: TN52D has inventory (600k)    │
│  • Real-time says: TN52D is YOUR_LOCATION        │
│  ✓ Recommendation: Use TN52D (consolidation)     │
│                                                  │
│  Speed: ~5ms 🎯                                  │
└──────────────────────────────────────────────────┘

Total Response Time: < 100ms"""

p = doc.add_paragraph(flow_diagram)
p.paragraph_format.left_indent = Inches(0.1)
p.runs[0].font.name = 'Courier New'
p.runs[0].font.size = Pt(8.5)

doc.add_paragraph()

doc.add_heading('Benefits of Hybrid Approach', 2)

benefits = [
    ('Speed', 'Knowledge base loads once (98 MB), no repeated queries'),
    ('Accuracy', 'Real-time queries ensure current availability'),
    ('Intelligence', 'Learn from history, validate with reality'),
    ('Reliability', 'Always uses latest data for critical decisions')
]

for title, desc in benefits:
    p = doc.add_paragraph()
    run = p.add_run(f'{title}: ')
    run.bold = True
    run.font.color.rgb = RGBColor(0, 102, 204)
    run.font.size = Pt(12)
    p.add_run(desc)

doc.add_page_break()

# =============================================================================
# SECTION 9: STATISTICS
# =============================================================================

doc.add_heading('9. Pattern Quality & Statistics', 1)

doc.add_heading('Overall Learning Results', 2)

stats_table = """Input Data:
  • Total Transactions Analyzed:  224,081
  • Unique Parts in Database:     3,215
  • Unique Locations:              31,416
  • Unique Clients:                87

Output Data:
  • Learned Patterns Created:      2,492
  • Coverage:                       77.5% of all parts
  • Average Putaways per Part:     90 transactions
  • Strong Patterns (>10 uses):    1,847 parts
  • Weak Patterns (2-10 uses):     645 parts"""

p = doc.add_paragraph(stats_table)
p.paragraph_format.left_indent = Inches(0.25)
p.runs[0].font.name = 'Courier New'
p.runs[0].font.size = Pt(10)

doc.add_paragraph()

doc.add_heading('Pattern Quality Distribution', 2)

quality_data = [
    ('High Confidence', '>50% usage on top location', '892 parts (35.8%)'),
    ('Medium Confidence', '20-50% usage on top location', '1,156 parts (46.4%)'),
    ('Low Confidence', '<20% usage on top location', '444 parts (17.8%)')
]

for quality, criteria, count in quality_data:
    p = doc.add_paragraph()
    run = p.add_run(f'{quality}: ')
    run.bold = True
    if 'High' in quality:
        run.font.color.rgb = RGBColor(0, 128, 0)
    elif 'Medium' in quality:
        run.font.color.rgb = RGBColor(255, 165, 0)
    else:
        run.font.color.rgb = RGBColor(255, 0, 0)
    p.add_run(f'{criteria}\n  → {count}')

doc.add_paragraph()

doc.add_heading('Example Patterns', 3)

example_high = """High Confidence Example - Part 1234:
  • Total putaways: 150
  • Top location: TP49A (78 times = 52%)
  • Pattern: STRONG ✓
  • Recommendation confidence: 95%"""

p = doc.add_paragraph(example_high)
p.paragraph_format.left_indent = Inches(0.25)
p.runs[0].font.color.rgb = RGBColor(0, 128, 0)

doc.add_paragraph()

example_low = """Low Confidence Example - Part 5678:
  • Total putaways: 12
  • Top location: SG01J (3 times = 25%)
  • Pattern: WEAK (spread across many locations)
  • Recommendation confidence: 35%"""

p = doc.add_paragraph(example_low)
p.paragraph_format.left_indent = Inches(0.25)
p.runs[0].font.color.rgb = RGBColor(255, 165, 0)

doc.add_page_break()

# =============================================================================
# SECTION 10: RECOMMENDATION FLOW
# =============================================================================

doc.add_heading('10. Recommendation Flow', 1)

doc.add_paragraph('When a warehouse operator requests a recommendation:')
doc.add_paragraph()

flow_steps = """1. User enters Part ID (e.g., 600)
   ↓
2. System retrieves learned pattern from Qdrant
   • Historical locations: TN52D (28.3%), SG01J (15.1%)
   • Primary zone: T
   ↓
3. System queries current inventory (MySQL)
   • TN52D: 600,000 units
   • TP55C: 600,000 units
   ↓
4. System checks location availability (MySQL)
   • TN52D: clientId = 34 → YOUR_LOCATION ✓
   • TP55C: clientId = 34 → YOUR_LOCATION ✓
   ↓
5. System ranks recommendations
   • #1: TN52D (historical + consolidation)
   • #2: TP55C (consolidation)
   • #3: SG01J (historical, if free)
   ↓
6. AI generates explanation (Google Gemini)
   "Location TN52D is recommended as it has been
    the most frequently used location for this part
    historically. The location is currently FREE
    and ready for immediate use."
   ↓
7. Display to user with alternatives"""

p = doc.add_paragraph(flow_steps)
p.paragraph_format.left_indent = Inches(0.25)

doc.add_page_break()

# =============================================================================
# SECTION 11: COMPLETE EXAMPLE
# =============================================================================

doc.add_heading('11. Complete Example: Part 600 Journey', 1)

doc.add_heading('Scenario: Learning Pattern for Part 600 (Bearing)', 2)

doc.add_heading('Step 1: Collect Raw Transactions', 3)

transactions_sample = """Transaction History (53 total):
2024-01-10: Part 600 → TN52D (ABC Corp)
2024-01-15: Part 600 → SG01J (ABC Corp)
2024-01-20: Part 600 → TN52D (ABC Corp)
2024-02-05: Part 600 → TN52D (ABC Corp)
2024-02-12: Part 600 → TP03D (ABC Corp)
2024-03-01: Part 600 → TN52D (ABC Corp)
... (47 more transactions)"""

p = doc.add_paragraph(transactions_sample)
p.paragraph_format.left_indent = Inches(0.25)
p.runs[0].font.name = 'Courier New'
p.runs[0].font.size = Pt(9)

doc.add_paragraph()

doc.add_heading('Step 2: Run Aggregation Query', 3)

query_results = """MySQL Aggregation Results:

Location  | Count | Percentage | Pattern
----------|-------|------------|----------
TN52D     | 15    | 28.3%      | Primary ⭐
SG01J     | 8     | 15.1%      | Secondary
TP03D     | 3     | 5.66%      | Tertiary
TN43D     | 1     | 1.9%       | Rare
(+32 more locations with smaller counts)"""

p = doc.add_paragraph(query_results)
p.paragraph_format.left_indent = Inches(0.25)
p.runs[0].font.name = 'Courier New'
p.runs[0].font.size = Pt(9)

doc.add_paragraph()

doc.add_heading('Step 3: Create Pattern JSON', 3)

final_json = """{
  "part_id": 600,
  "part_code": "42645EQ",
  "total_putaways": 53,
  "primary_zone": "T",
  "all_locations": [
    {"code": "TN52D", "count": 15, "percentage": 28.3},
    {"code": "SG01J", "count": 8, "percentage": 15.1},
    {"code": "TP03D", "count": 3, "percentage": 5.66}
  ],
  "pattern_strength": "MEDIUM",
  "confidence_score": 0.72
}"""

p = doc.add_paragraph(final_json)
p.paragraph_format.left_indent = Inches(0.25)
p.runs[0].font.name = 'Courier New'
p.runs[0].font.size = Pt(9)

doc.add_paragraph()

doc.add_heading('Step 4: Store in Qdrant & Use for Recommendations', 3)

p = doc.add_paragraph()
run = p.add_run('Pattern saved to PartSummary collection\n')
run.bold = True
run = p.add_run('Status: Ready for recommendations ✓\n')
run.font.color.rgb = RGBColor(0, 128, 0)
run.bold = True
p.add_run('\nNow whenever Part 600 arrives, the system will recommend TN52D as the primary location (28.3% historical usage).')

doc.add_page_break()

# =============================================================================
# SECTION 12: BENEFITS
# =============================================================================

doc.add_heading('12. Benefits & Advantages', 1)

doc.add_heading('1. Data-Driven Decisions', 2)
benefits1 = [
    'Based on 224,081 real historical transactions',
    'Not guessing - using proven patterns',
    'Reduces human error in location selection',
    'Eliminates need to memorize part-location mappings'
]
for b in benefits1:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('2. Speed & Efficiency', 2)
benefits2 = [
    'Recommendations in < 100ms',
    'Pre-computed patterns (no waiting)',
    'Fast Qdrant lookups (vector database)',
    'Minimal real-time queries (only what\'s needed)'
]
for b in benefits2:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('3. Adaptability & Learning', 2)
benefits3 = [
    'Patterns update with new transactions',
    'System learns from changing practices',
    'Confidence scores adjust automatically',
    'Improves accuracy over time'
]
for b in benefits3:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('4. Transparency & Trust', 2)
benefits4 = [
    'Shows WHY each location is recommended',
    'Provides usage statistics (15×, 28.3%)',
    'Displays alternative options',
    'Allows manual overrides with audit logging'
]
for b in benefits4:
    doc.add_paragraph(b, style='List Bullet')

doc.add_page_break()

# =============================================================================
# SECTION 13: UPDATE SCHEDULE
# =============================================================================

doc.add_heading('13. Update Schedule & Maintenance', 1)

doc.add_heading('Daily Updates (Automated)', 2)

daily = """# Recommended: Nightly cron job
python add_inventory_patterns.py
python upload_to_qdrant.py

Updates:
  • Inventory snapshots
  • Recent transaction patterns
  • New part additions"""

p = doc.add_paragraph(daily)
p.paragraph_format.left_indent = Inches(0.25)
p.runs[0].font.name = 'Courier New'
p.runs[0].font.size = Pt(9)

doc.add_paragraph()

doc.add_heading('Weekly Review (Manual)', 2)

weekly = [
    'Review pattern changes and new high-volume parts',
    'Check confidence scores for accuracy',
    'Verify no data quality issues',
    'Monitor system performance metrics'
]
for item in weekly:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()

doc.add_heading('Monthly Maintenance (Comprehensive)', 2)

monthly = [
    'Full knowledge base regeneration from scratch',
    'Statistical analysis of pattern changes',
    'Client preference updates',
    'System performance audit and optimization',
    'Documentation updates if needed'
]
for item in monthly:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()

doc.add_heading('Data Freshness Impact', 2)

p = doc.add_paragraph()
run = p.add_run('Important Note: ')
run.bold = True
run.font.color.rgb = RGBColor(255, 0, 0)
p.add_run('Even if the knowledge base is outdated, the system continues to work correctly because real-time MySQL queries always provide current inventory and availability data. Historical patterns are used as guidance, not absolute rules.')

doc.add_page_break()

# =============================================================================
# CONCLUSION & FOOTER
# =============================================================================

doc.add_heading('Conclusion', 1)

conclusion_text = doc.add_paragraph()
conclusion_text.add_run('The ').font.size = Pt(11)
run = conclusion_text.add_run('StockRight Pattern Learning System')
run.bold = True
run.font.color.rgb = RGBColor(0, 102, 204)
conclusion_text.add_run(' represents a sophisticated approach to warehouse management, combining:')

doc.add_paragraph()

conclusion_points = [
    ('Historical Intelligence', 'Learns from 224,081 transactions'),
    ('SQL Aggregation', 'Extracts patterns using 6-step pipeline'),
    ('Vector Storage', 'Stores 2,492 patterns in Qdrant for fast access'),
    ('Hybrid Architecture', 'Combines offline patterns with real-time data'),
    ('AI Explanations', 'Provides natural language recommendations via Google Gemini'),
    ('Production-Ready', 'Fast (<100ms), accurate, and transparent')
]

for title, desc in conclusion_points:
    p = doc.add_paragraph()
    run = p.add_run(f'✓ {title}: ')
    run.bold = True
    run.font.color.rgb = RGBColor(0, 128, 0)
    p.add_run(desc)

doc.add_paragraph()

final_note = doc.add_paragraph()
final_note.add_run('This data-driven approach ensures ').font.size = Pt(11)
run = final_note.add_run('efficient warehouse operations')
run.bold = True
final_note.add_run(' while maintaining ')
run = final_note.add_run('flexibility for manual overrides')
run.bold = True
final_note.add_run(' when needed.')

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

# Final footer
footer_para = doc.add_paragraph()
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_para.add_run('━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n\n').font.color.rgb = RGBColor(128, 128, 128)
run = footer_para.add_run('StockRight Agentic Logistics Engine (SALE)\n')
run.font.size = Pt(11)
run.bold = True
run.font.color.rgb = RGBColor(0, 102, 204)
footer_para.add_run('Document Version: 1.0 | February 2026\n').font.size = Pt(9)
footer_para.add_run('Knowledge Base: 224,081 Transactions | 2,492 Learned Patterns').font.size = Pt(9)

# Add page numbers
add_page_number(doc)

# Save document
output_path = 'C:\\Users\\lenovo\\warehouse-qdrant-system\\StockRight_Complete_Documentation.docx'
doc.save(output_path)
print("Professional document created successfully!")
print(f"Location: {output_path}")
print("\nDocument includes:")
print("  - Cover page with branding")
print("  - Table of contents")
print("  - 13 comprehensive sections")
print("  - Code examples and diagrams")
print("  - Color-coded content")
print("  - Page numbers")
print("  - Professional formatting")
